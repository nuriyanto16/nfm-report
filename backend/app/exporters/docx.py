"""Render laporan ke Word (.docx). Sumber tunggal untuk PDF juga (lihat pdf.py)."""
from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.core.report import ReportResult
from app.exporters.layout import COLUMNS, HEADERS, row_values

_HEADER_BG = "1F4E78"
_STATUS_BG = {
    "Done": "C6EFCE", "Progress": "FFEB9C", "Hold": "FFC7CE",
    "Back Log": "D9D9D9", "To Do": "BDD7EE",
}


def _shade(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color,
    })
    tcPr.append(shd)


def _set_cell(cell, text, *, bold=False, color=None, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def render_docx(result: ReportResult) -> bytes:
    doc = Document()
    section = doc.sections[0]
    # Landscape A4 agar tabel lebar muat.
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, m, Pt(36))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FAST REPORT")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(_HEADER_BG)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(result.source_name)
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor.from_string("555555")

    p = result.period
    scope = {"daily": "Harian", "weekly": "Mingguan",
             "monthly": "Bulanan"}.get(p.mode, p.mode)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Laporan {scope} · {p.label}  (acuan tanggal: {p.date_field})\n"
        f"Total: {len(result.rows)} task   ·   "
        f"Dibuat: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ).font.size = Pt(9)

    _add_table(doc, result)
    doc.add_paragraph()
    _add_summary(doc, result)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_table(doc, result: ReportResult):
    table = doc.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(HEADERS):
        cell = table.rows[0].cells[i]
        _set_cell(cell, header, bold=True, color="FFFFFF", size=8)
        _shade(cell, _HEADER_BG)
    # Ulangi baris header di tiap halaman.
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(trPr.makeelement(qn("w:tblHeader"), {qn("w:val"): "true"}))

    status_idx = HEADERS.index("Status")
    for rec in result.rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values(rec)):
            _set_cell(cells[i], value, size=8)
        bg = _STATUS_BG.get(rec.status)
        if bg:
            _shade(cells[status_idx], bg)


def _add_summary(doc, result: ReportResult):
    s = result.summary
    doc.add_paragraph().add_run("RINGKASAN").bold = True
    for title, data in (
        ("Per Status", s["by_status"]),
        ("Per Aplikasi", s["by_aplikasi"]),
        ("Per PIC", s["by_pic"]),
    ):
        doc.add_paragraph().add_run(title).bold = True
        line = ", ".join(f"{k}: {v}" for k, v in data.items())
        doc.add_paragraph(line).runs[0].font.size = Pt(9)
